import os
import mimetypes
import json
import tempfile
import shutil
from flask import render_template, send_file, abort, Response, jsonify, request, redirect, url_for, flash
from flask_login import login_required, current_user
from . import forms_bp
from app import db
from models import FormPreset
import openpyxl
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import docx

# Diretório base dos formulários
FORMS_DIR = os.path.join(os.getcwd(), 'extracted_forms')

@forms_bp.route('/')
@login_required
def index():
    """Página principal para visualização de formulários."""
    # Obtém as categorias de formulários (pastas)
    categories = []
    for item in os.listdir(FORMS_DIR):
        if os.path.isdir(os.path.join(FORMS_DIR, item)):
            categories.append(item)
    
    return render_template(
        'forms/index.html',
        title="Formulários Zelopack",
        categories=categories
    )

@forms_bp.route('/category/<category>')
@login_required
def category(category):
    """Visualizar formulários de uma categoria específica."""
    category_path = os.path.join(FORMS_DIR, category)
    
    # Verificar se a categoria existe
    if not os.path.exists(category_path) or not os.path.isdir(category_path):
        abort(404)
    
    # Listar formulários na categoria
    forms = []
    for root, dirs, files in os.walk(category_path):
        relative_path = os.path.relpath(root, FORMS_DIR)
        
        for file in files:
            if file.startswith('~$') or file.startswith('.'):  # Ignorar arquivos temporários
                continue
                
            file_path = os.path.join(relative_path, file)
            file_ext = os.path.splitext(file)[1].lower()
            
            # Ícone baseado na extensão
            icon = 'fa-file'
            if file_ext in ['.pdf']:
                icon = 'fa-file-pdf'
            elif file_ext in ['.doc', '.docx']:
                icon = 'fa-file-word'
            elif file_ext in ['.xls', '.xlsx']:
                icon = 'fa-file-excel'
            elif file_ext in ['.ppt', '.pptx']:
                icon = 'fa-file-powerpoint'
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif']:
                icon = 'fa-file-image'
            
            forms.append({
                'name': file,
                'path': file_path,
                'icon': icon,
                'date': os.path.getmtime(os.path.join(FORMS_DIR, file_path))
            })
    
    # Ordenar formulários por nome
    forms.sort(key=lambda x: x['name'])
    
    return render_template(
        'forms/category.html',
        title=f"Formulários - {category}",
        category=category,
        forms=forms
    )

@forms_bp.route('/view/<path:file_path>')
@login_required
def view_form(file_path):
    """Visualizar um formulário específico."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    # Obter o tipo MIME do arquivo
    mime_type, _ = mimetypes.guess_type(full_path)
    if not mime_type:
        mime_type = 'application/octet-stream'
    
    # Enviar o arquivo para visualização
    return send_file(
        full_path,
        mimetype=mime_type,
        as_attachment=False,
        download_name=os.path.basename(full_path)
    )

@forms_bp.route('/download/<path:file_path>')
@login_required
def download_form(file_path):
    """Baixar um formulário específico."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    # Enviar o arquivo para download
    return send_file(
        full_path,
        as_attachment=True,
        download_name=os.path.basename(full_path)
    )

@forms_bp.route('/search')
@login_required
def search_forms():
    """Pesquisar formulários."""
    query = request.args.get('q', '').lower()
    if not query or len(query) < 3:
        return jsonify([])
    
    results = []
    
    # Buscar em todas as pastas
    for root, dirs, files in os.walk(FORMS_DIR):
        for file in files:
            if file.startswith('~$') or file.startswith('.'):  # Ignorar arquivos temporários
                continue
                
            if query in file.lower():
                relative_path = os.path.relpath(root, FORMS_DIR)
                file_path = os.path.join(relative_path, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                # Ícone baseado na extensão
                icon = 'fa-file'
                if file_ext in ['.pdf']:
                    icon = 'fa-file-pdf'
                elif file_ext in ['.doc', '.docx']:
                    icon = 'fa-file-word'
                elif file_ext in ['.xls', '.xlsx']:
                    icon = 'fa-file-excel'
                elif file_ext in ['.ppt', '.pptx']:
                    icon = 'fa-file-powerpoint'
                elif file_ext in ['.jpg', '.jpeg', '.png', '.gif']:
                    icon = 'fa-file-image'
                
                results.append({
                    'name': file,
                    'path': file_path,
                    'category': os.path.basename(root),
                    'icon': icon
                })
    
    # Ordenar resultados por nome
    results.sort(key=lambda x: x['name'])
    
    return jsonify(results)


@forms_bp.route('/api/preset/<int:preset_id>')
@login_required
def get_preset_data(preset_id):
    """API para obter dados de uma predefinição."""
    preset = FormPreset.query.get_or_404(preset_id)
    
    # Verificar se o usuário tem permissão (qualquer usuário pode ver)
    return jsonify({
        'success': True,
        'preset': {
            'id': preset.id,
            'name': preset.name,
            'description': preset.description,
            'form_type': preset.form_type,
            'data': preset.data,
            'is_default': preset.is_default
        }
    })


@forms_bp.route('/fill/<path:file_path>')
@login_required
def fill_form(file_path):
    """Interface para preencher um formulário."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    file_name = os.path.basename(full_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    # Obter campos disponíveis para preenchimento (baseado no tipo de arquivo)
    fields = get_form_fields(full_path)
    
    # Obter predefinições disponíveis para esse formulário
    presets = FormPreset.query.filter_by(
        form_type=file_name,
        is_active=True
    ).order_by(FormPreset.is_default.desc(), FormPreset.name).all()
    
    return render_template(
        'forms/fill_form.html',
        title=f"Preencher Formulário - {file_name}",
        file_path=file_path,
        file_name=file_name,
        file_ext=file_ext,
        fields=fields,
        presets=presets
    )


@forms_bp.route('/presets/<path:file_path>', methods=['GET'])
@login_required
def list_presets(file_path):
    """Lista todas as predefinições para um formulário específico."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    file_name = os.path.basename(full_path)
    
    # Obter predefinições disponíveis para esse formulário
    presets = FormPreset.query.filter_by(
        form_type=file_name,
        is_active=True
    ).order_by(FormPreset.is_default.desc(), FormPreset.name).all()
    
    return render_template(
        'forms/presets.html',
        title=f"Predefinições - {file_name}",
        file_path=file_path,
        file_name=file_name,
        presets=[preset.to_dict() for preset in presets]
    )


@forms_bp.route('/presets/create/<path:file_path>', methods=['GET', 'POST'])
@login_required
def create_preset(file_path):
    """Criar uma nova predefinição para um formulário."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    file_name = os.path.basename(full_path)
    
    if request.method == 'POST':
        # Obter dados do formulário
        data = request.form.to_dict()
        fields_data = {}
        
        # Verificar campos especiais (não são campos do formulário)
        preset_name = data.pop('preset_name', f'Predefinição para {file_name}')
        preset_description = data.pop('preset_description', '')
        is_default = data.pop('is_default', 'off') == 'on'
        
        # Os demais dados são campos do formulário
        for key, value in data.items():
            if key.startswith('field_'):
                field_id = key.replace('field_', '')
                fields_data[field_id] = value
        
        # Criar uma nova predefinição
        preset = FormPreset(
            name=preset_name,
            description=preset_description,
            form_type=file_name,
            file_path=file_path,
            created_by=current_user.id,
            data=fields_data,
            is_default=is_default
        )
        
        # Se esta predefinição é marcada como padrão, desmarcar as demais
        if is_default:
            FormPreset.query.filter_by(
                form_type=file_name,
                is_default=True
            ).update({'is_default': False})
        
        db.session.add(preset)
        db.session.commit()
        
        flash(f'Predefinição "{preset_name}" criada com sucesso!', 'success')
        return redirect(url_for('forms.list_presets', file_path=file_path))
    
    # Obter campos disponíveis para preenchimento (baseado no tipo de arquivo)
    fields = get_form_fields(full_path)
    
    return render_template(
        'forms/create_preset.html',
        title=f"Nova Predefinição - {file_name}",
        file_path=file_path,
        file_name=file_name,
        fields=fields
    )


@forms_bp.route('/presets/edit/<int:preset_id>', methods=['GET', 'POST'])
@login_required
def edit_preset(preset_id):
    """Editar uma predefinição existente."""
    preset = FormPreset.query.get_or_404(preset_id)
    
    # Verificar se o usuário tem permissão (criador ou administrador)
    if preset.created_by != current_user.id and not current_user.is_admin:
        abort(403)
    
    full_path = os.path.join(FORMS_DIR, preset.file_path)
    
    # Verificar se o arquivo ainda existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('O arquivo original não está mais disponível.', 'warning')
        return redirect(url_for('forms.list_presets', file_path=preset.file_path))
    
    if request.method == 'POST':
        # Obter dados do formulário
        data = request.form.to_dict()
        fields_data = {}
        
        # Verificar campos especiais (não são campos do formulário)
        preset.name = data.pop('preset_name', preset.name)
        preset.description = data.pop('preset_description', preset.description)
        is_default = data.pop('is_default', 'off') == 'on'
        
        # Os demais dados são campos do formulário
        for key, value in data.items():
            if key.startswith('field_'):
                field_id = key.replace('field_', '')
                fields_data[field_id] = value
        
        preset.data = fields_data
        
        # Atualizar status padrão
        if is_default and not preset.is_default:
            # Se esta predefinição agora é marcada como padrão, desmarcar as demais
            FormPreset.query.filter_by(
                form_type=preset.form_type,
                is_default=True
            ).update({'is_default': False})
            preset.is_default = True
        elif not is_default and preset.is_default:
            preset.is_default = False
        
        db.session.commit()
        
        flash(f'Predefinição "{preset.name}" atualizada com sucesso!', 'success')
        return redirect(url_for('forms.list_presets', file_path=preset.file_path))
    
    # Obter campos disponíveis para preenchimento (baseado no tipo de arquivo)
    fields = get_form_fields(full_path)
    
    return render_template(
        'forms/edit_preset.html',
        title=f"Editar Predefinição - {preset.name}",
        preset=preset,
        file_path=preset.file_path,
        file_name=os.path.basename(full_path),
        fields=fields
    )


@forms_bp.route('/presets/delete/<int:preset_id>', methods=['POST'])
@login_required
def delete_preset(preset_id):
    """Excluir uma predefinição."""
    preset = FormPreset.query.get_or_404(preset_id)
    
    # Verificar se o usuário tem permissão (criador ou administrador)
    if preset.created_by != current_user.id and not current_user.is_admin:
        abort(403)
    
    file_path = preset.file_path
    preset_name = preset.name
    
    db.session.delete(preset)
    db.session.commit()
    
    flash(f'Predefinição "{preset_name}" excluída com sucesso!', 'success')
    return redirect(url_for('forms.list_presets', file_path=file_path))


@forms_bp.route('/download_filled/<path:file_path>', methods=['POST'])
@login_required
def download_filled_form(file_path):
    """Baixar um formulário preenchido com os dados fornecidos."""
    full_path = os.path.join(FORMS_DIR, file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        abort(404)
    
    # Obter os dados para preenchimento
    form_data = {}
    
    # Verificar se está usando uma predefinição ou dados do formulário
    if 'preset_id' in request.form and request.form['preset_id']:
        preset_id = request.form['preset_id']
        preset = FormPreset.query.get_or_404(preset_id)
        form_data = preset.data
    else:
        # Dados enviados diretamente pelo formulário
        for key, value in request.form.items():
            if key.startswith('field_'):
                field_id = key.replace('field_', '')
                form_data[field_id] = value
    
    # Gerar arquivo preenchido
    file_name = os.path.basename(full_path)
    file_ext = os.path.splitext(file_name)[1].lower()
    
    filled_path = fill_form_with_data(full_path, form_data)
    
    if not filled_path:
        flash('Não foi possível preencher o formulário. Formato não suportado.', 'error')
        return redirect(url_for('forms.fill_form', file_path=file_path))
    
    # Enviar o arquivo preenchido para download
    return send_file(
        filled_path,
        as_attachment=True,
        download_name=f'Preenchido_{file_name}'
    )


@forms_bp.route('/preset/<int:preset_id>/download')
@login_required
def download_preset(preset_id):
    """Baixar um formulário preenchido com uma predefinição específica."""
    preset = FormPreset.query.get_or_404(preset_id)
    full_path = os.path.join(FORMS_DIR, preset.file_path)
    
    # Verificar se o arquivo existe
    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        flash('O arquivo original não está mais disponível.', 'warning')
        return redirect(url_for('forms.index'))
    
    filled_path = fill_form_with_data(full_path, preset.data)
    
    if not filled_path:
        flash('Não foi possível preencher o formulário. Formato não suportado.', 'error')
        return redirect(url_for('forms.list_presets', file_path=preset.file_path))
    
    file_name = os.path.basename(full_path)
    
    # Enviar o arquivo preenchido para download
    return send_file(
        filled_path,
        as_attachment=True,
        download_name=f'{preset.name}_{file_name}'
    )


# Funções auxiliares para manipulação de formulários

def get_form_fields(file_path):
    """Extrai os campos disponíveis para preenchimento em um formulário."""
    file_ext = os.path.splitext(file_path)[1].lower()
    fields = []
    
    try:
        if file_ext == '.xlsx' or file_ext == '.xls':
            # Extrair campos de planilha Excel
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            sheet = workbook.active
            
            if sheet is None:
                flash('Não foi possível acessar a planilha.', 'warning')
                return fields
                
            # Obter o número máximo de linhas e colunas com segurança
            max_row = sheet.max_row or 1
            max_col = sheet.max_column or 1
            
            for row in range(1, max_row + 1):
                for col in range(1, max_col + 1):
                    cell = sheet.cell(row=row, column=col)
                    if cell is None:
                        continue
                    
                    value = cell.value
                    
                    if value and isinstance(value, str) and ('___' in value or '____' in value):
                        # Encontrou um campo para preenchimento (representado por sublinhados)
                        field_id = f"cell_{row}_{col}"
                        sheet_title = getattr(sheet, 'title', 'Planilha')
                        fields.append({
                            'id': field_id,
                            'name': f"Campo em {sheet_title} ({get_column_letter(col)}{row})",
                            'value': ''
                        })
            
        elif file_ext == '.docx':
            # Extrair campos de documento Word
            doc = docx.Document(file_path)
            field_index = 0
            
            for para_index, para in enumerate(doc.paragraphs):
                text = para.text
                if '___' in text or '____' in text:
                    # Encontrou um parágrafo com campos para preenchimento
                    field_id = f"para_{para_index}"
                    fields.append({
                        'id': field_id,
                        'name': f"Campo no parágrafo {para_index + 1}: {text[:30]}...",
                        'value': ''
                    })
            
        elif file_ext == '.pdf':
            # Extrair campos de PDF (mais complexo)
            # Implementação básica para detecção de campos
            reader = PyPDF2.PdfReader(file_path)
            form_fields = reader.get_fields()
            
            if form_fields:
                # PDF tem campos de formulário
                for field_name, field_value in form_fields.items():
                    fields.append({
                        'id': field_name,
                        'name': field_name,
                        'value': ''
                    })
            else:
                # PDF não tem campos de formulário, tentar identificar por texto
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text and ('___' in text or '____' in text):
                        # Encontrou texto com campos potenciais
                        field_id = f"pdf_page_{page_num}"
                        fields.append({
                            'id': field_id,
                            'name': f"Campo na página {page_num + 1}",
                            'value': ''
                        })
    except Exception as e:
        print(f"Erro ao extrair campos do formulário: {e}")
    
    return fields


def fill_form_with_data(file_path, form_data):
    """Preenche um formulário com os dados fornecidos e retorna o caminho do arquivo preenchido."""
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Criar diretório temporário para o formulário preenchido
    temp_dir = tempfile.mkdtemp()
    file_name = os.path.basename(file_path)
    output_path = os.path.join(temp_dir, f'filled_{file_name}')
    
    try:
        if file_ext == '.xlsx' or file_ext == '.xls':
            # Preencher planilha Excel
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active
            
            if sheet is None:
                raise ValueError('Não foi possível acessar a planilha.')
            
            # Aplicar os dados aos campos identificados
            for field_id, value in form_data.items():
                if field_id.startswith('cell_'):
                    _, row, col = field_id.split('_')
                    row, col = int(row), int(col)
                    
                    try:
                        cell = sheet.cell(row=row, column=col)
                        if cell is None:
                            continue
                            
                        cell.value = value
                        
                        # Aplicar estilo para destacar o campo preenchido
                        cell.font = Font(bold=True, color="0000FF")
                    except Exception as e:
                        print(f"Erro ao preencher célula ({row},{col}): {e}")
            
            # Salvar a planilha preenchida
            workbook.save(output_path)
            
        elif file_ext == '.docx':
            # Preencher documento Word
            doc = docx.Document(file_path)
            
            # Aplicar os dados aos campos identificados
            for field_id, value in form_data.items():
                if field_id.startswith('para_'):
                    para_index = int(field_id.split('_')[1])
                    if para_index < len(doc.paragraphs):
                        para = doc.paragraphs[para_index]
                        text = para.text
                        
                        # Substituir campos em branco pelo valor
                        new_text = text.replace('_____', value).replace('____', value).replace('___', value)
                        
                        # Limpar o parágrafo e adicionar o texto substituído
                        para.clear()
                        para.add_run(new_text)
            
            # Salvar o documento preenchido
            doc.save(output_path)
            
        elif file_ext == '.pdf':
            # Para PDFs, criamos um novo PDF com o texto sobreposto
            # (abordagem simples, funcionalidade limitada)
            reader = PyPDF2.PdfReader(file_path)
            writer = PyPDF2.PdfWriter()
            
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                writer.add_page(page)
            
            # Se o PDF tem campos de formulário
            form_fields = reader.get_fields()
            if form_fields:
                # Preencher campos de formulário
                for field_id, value in form_data.items():
                    if field_id in form_fields:
                        writer.update_page_form_field_values(0, {field_id: value})
            
            # Salvar o PDF preenchido
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
                
        else:
            # Formato não suportado
            shutil.rmtree(temp_dir)
            return None
    
    except Exception as e:
        print(f"Erro ao preencher formulário: {e}")
        shutil.rmtree(temp_dir)
        return None
    
    return output_path


# A função get_column_letter já está importada de openpyxl.utils