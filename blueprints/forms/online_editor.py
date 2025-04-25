"""
Módulo de Editor Online Universal para ZeloPack
Este módulo oferece funções para converter e editar arquivos online
independente do formato original
"""

import os
import json
import tempfile
import uuid
import base64
import io
from datetime import datetime
from pathlib import Path

import openpyxl
import docx
import PyPDF2
import pandas as pd
from flask import current_app
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4


class OnlineEditorConverter:
    """Converte diferentes tipos de documentos para formato editável online e vice-versa"""
    
    def __init__(self, file_path, forms_dir=None):
        """
        Inicializa o conversor com o caminho do arquivo
        
        Args:
            file_path: Caminho relativo ao FORMS_DIR
            forms_dir: Diretório base dos formulários (opcional)
        """
        self.forms_dir = forms_dir or os.environ.get('FORMS_DIR', 'attached_assets')
        self.file_path = file_path
        self.full_path = os.path.join(self.forms_dir, file_path)
        self.file_name = os.path.basename(self.full_path)
        self.file_ext = os.path.splitext(self.file_name)[1].lower()
        self.temp_dir = None
        
    def get_editable_content(self):
        """
        Extrai o conteúdo do arquivo e retorna uma estrutura de dados editável
        
        Returns:
            dict: Estrutura com dados editáveis e metadados
        """
        try:
            if self.file_ext in ['.xlsx', '.xls']:
                return self._extract_excel_content()
            elif self.file_ext == '.docx':
                return self._extract_docx_content()
            elif self.file_ext == '.pdf':
                return self._extract_pdf_content()
            else:
                return {
                    'success': False,
                    'message': f'Formato não suportado: {self.file_ext}',
                    'content_type': 'unsupported'
                }
        except Exception as e:
            print(f"Erro ao converter arquivo: {e}")
            return {
                'success': False,
                'message': f'Erro ao processar o arquivo: {str(e)}',
                'content_type': 'error'
            }
    
    def _extract_excel_content(self):
        """Extrai conteúdo de arquivos Excel"""
        try:
            # Ler o Excel com pandas para facilitar a manipulação
            excel_data = {}
            
            # Ler todas as abas do Excel
            xls = pd.ExcelFile(self.full_path)
            sheet_names = xls.sheet_names
            
            # Criar um dicionário com os dados de cada aba
            for sheet_name in sheet_names:
                df = pd.read_excel(self.full_path, sheet_name=sheet_name)
                
                # Converter para o formato de células editáveis
                sheet_data = []
                for i, row in df.iterrows():
                    row_data = []
                    for col_name in df.columns:
                        cell_value = row[col_name]
                        # Converter valores para string se necessário
                        if pd.isna(cell_value):
                            cell_str = ""
                        else:
                            # Formatar valores especiais
                            if isinstance(cell_value, (datetime, pd.Timestamp)):
                                cell_str = cell_value.strftime('%d/%m/%Y')
                            else:
                                cell_str = str(cell_value)
                        
                        # Identificar células que parecem ser campos para preenchimento
                        is_field = False
                        if isinstance(cell_str, str) and ('___' in cell_str or '____' in cell_str):
                            is_field = True
                        
                        row_data.append({
                            'value': cell_str,
                            'is_field': is_field,
                            'original_type': str(type(cell_value).__name__),
                            'row': i + 1,  # +1 para corresponder ao índice de linha do Excel
                            'col': df.columns.get_loc(col_name) + 1  # +1 para corresponder ao índice de coluna do Excel
                        })
                    sheet_data.append(row_data)
                
                excel_data[sheet_name] = {
                    'data': sheet_data,
                    'columns': df.columns.tolist(),
                    'row_count': len(df),
                    'col_count': len(df.columns)
                }
            
            # Metadados do arquivo
            workbook = openpyxl.load_workbook(self.full_path, data_only=True)
            active_sheet = workbook.active.title if workbook.active else sheet_names[0]
            
            return {
                'success': True,
                'content_type': 'excel',
                'file_name': self.file_name,
                'sheets': excel_data,
                'active_sheet': active_sheet,
                'total_sheets': len(sheet_names),
                'fields': self._get_excel_fields(workbook)
            }
            
        except Exception as e:
            print(f"Erro ao processar Excel: {e}")
            return {
                'success': False,
                'message': f'Erro ao processar planilha: {str(e)}',
                'content_type': 'error'
            }
    
    def _extract_docx_content(self):
        """Extrai conteúdo de arquivos Word"""
        try:
            doc = docx.Document(self.full_path)
            paragraphs = []
            tables = []
            fields = []
            
            # Extrair parágrafos
            for i, para in enumerate(doc.paragraphs):
                text = para.text
                is_field = False
                
                # Identificar campos para preenchimento
                if '___' in text or '____' in text:
                    is_field = True
                    fields.append({
                        'id': f"para_{i}",
                        'name': f"Campo no parágrafo {i + 1}",
                        'value': text,
                        'is_field': True,
                        'paragraph_index': i
                    })
                
                paragraphs.append({
                    'text': text,
                    'is_field': is_field,
                    'paragraph_index': i
                })
            
            # Extrair tabelas
            for t, table in enumerate(doc.tables):
                table_data = []
                for r, row in enumerate(table.rows):
                    row_data = []
                    for c, cell in enumerate(row.cells):
                        text = cell.text
                        is_field = False
                        
                        # Identificar campos para preenchimento
                        if '___' in text or '____' in text:
                            is_field = True
                            fields.append({
                                'id': f"table_{t}_cell_{r}_{c}",
                                'name': f"Campo na tabela {t+1}, linha {r+1}, coluna {c+1}",
                                'value': text,
                                'is_field': True,
                                'table_index': t,
                                'row_index': r,
                                'col_index': c
                            })
                        
                        row_data.append({
                            'text': text,
                            'is_field': is_field,
                            'row_index': r,
                            'col_index': c
                        })
                    table_data.append(row_data)
                
                tables.append({
                    'table_index': t,
                    'row_count': len(table.rows),
                    'col_count': len(table.rows[0].cells) if table.rows else 0,
                    'data': table_data
                })
            
            return {
                'success': True,
                'content_type': 'docx',
                'file_name': self.file_name,
                'paragraphs': paragraphs,
                'tables': tables,
                'fields': fields,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables),
                'total_fields': len(fields)
            }
            
        except Exception as e:
            print(f"Erro ao processar Word: {e}")
            return {
                'success': False,
                'message': f'Erro ao processar documento Word: {str(e)}',
                'content_type': 'error'
            }
    
    def _extract_pdf_content(self):
        """Extrai conteúdo de arquivos PDF"""
        try:
            reader = PyPDF2.PdfReader(self.full_path)
            num_pages = len(reader.pages)
            pages_content = []
            fields = []
            
            # Extrair campos de formulário, se existirem
            form_fields = reader.get_fields()
            if form_fields:
                for field_name, field_value in form_fields.items():
                    fields.append({
                        'id': field_name,
                        'name': field_name,
                        'value': str(field_value if field_value else ''),
                        'is_form_field': True
                    })
            
            # Extrair texto de cada página
            for i in range(num_pages):
                page = reader.pages[i]
                text = page.extract_text()
                
                # Se não tem campos de formulário, tentar identificar campos por texto
                if not form_fields and ('___' in text or '____' in text):
                    fields.append({
                        'id': f"page_{i}_text",
                        'name': f"Campo na página {i+1}",
                        'value': text,
                        'page_index': i
                    })
                
                pages_content.append({
                    'page_index': i,
                    'text': text
                })
            
            # Gerar base64 do PDF para visualização inline
            with open(self.full_path, 'rb') as f:
                pdf_base64 = base64.b64encode(f.read()).decode('utf-8')
            
            return {
                'success': True,
                'content_type': 'pdf',
                'file_name': self.file_name,
                'pages': pages_content,
                'fields': fields,
                'total_pages': num_pages,
                'has_form_fields': bool(form_fields),
                'pdf_base64': pdf_base64,
                'original_path': self.file_path
            }
            
        except Exception as e:
            print(f"Erro ao processar PDF: {e}")
            return {
                'success': False,
                'message': f'Erro ao processar arquivo PDF: {str(e)}',
                'content_type': 'error'
            }
    
    def _get_excel_fields(self, workbook):
        """
        Identifica campos para preenchimento no Excel
        
        Args:
            workbook: Objeto Workbook do openpyxl
            
        Returns:
            list: Lista de campos identificados
        """
        fields = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            for row in range(1, sheet.max_row + 1):
                for col in range(1, sheet.max_column + 1):
                    cell = sheet.cell(row=row, column=col)
                    value = cell.value
                    
                    if value and isinstance(value, str) and ('___' in value or '____' in value):
                        col_letter = openpyxl.utils.get_column_letter(col)
                        
                        fields.append({
                            'id': f"sheet_{sheet_name}_cell_{row}_{col}",
                            'name': f"Campo em {sheet_name} ({col_letter}{row})",
                            'value': value,
                            'is_field': True,
                            'sheet': sheet_name,
                            'row': row,
                            'col': col,
                            'col_letter': col_letter
                        })
        
        return fields
    
    def save_edited_content(self, edited_data):
        """
        Salva o conteúdo editado de volta no formato original
        
        Args:
            edited_data (dict): Estrutura de dados com o conteúdo editado
            
        Returns:
            dict: Resultado da operação com caminho do arquivo salvo
        """
        try:
            # Criar diretório temporário para o arquivo salvo
            self.temp_dir = tempfile.mkdtemp(prefix='zelopack_edited_')
            
            # Gerar nome de arquivo com timestamp para evitar sobreposição
            timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
            filename_base = os.path.splitext(self.file_name)[0]
            new_filename = f"{filename_base}_preenchido_{timestamp}{self.file_ext}"
            output_path = os.path.join(self.temp_dir, new_filename)
            
            # Processar conforme o tipo de arquivo
            if self.file_ext in ['.xlsx', '.xls']:
                self._save_excel_content(edited_data, output_path)
            elif self.file_ext == '.docx':
                self._save_docx_content(edited_data, output_path)
            elif self.file_ext == '.pdf':
                self._save_pdf_content(edited_data, output_path)
            else:
                return {
                    'success': False,
                    'message': f'Formato não suportado para edição: {self.file_ext}'
                }
            
            return {
                'success': True,
                'message': 'Arquivo salvo com sucesso',
                'file_path': output_path,
                'file_name': new_filename
            }
            
        except Exception as e:
            print(f"Erro ao salvar conteúdo editado: {e}")
            return {
                'success': False,
                'message': f'Erro ao salvar o arquivo: {str(e)}'
            }
    
    def _save_excel_content(self, edited_data, output_path):
        """
        Salva conteúdo editado em formato Excel
        
        Args:
            edited_data (dict): Dados editados
            output_path (str): Caminho do arquivo de saída
        """
        # Copiar o arquivo original como base
        workbook = openpyxl.load_workbook(self.full_path)
        
        # Atualizar células com os valores editados
        for field in edited_data.get('fields', []):
            if not field.get('is_field'):
                continue
                
            sheet_name = field.get('sheet')
            row = field.get('row')
            col = field.get('col')
            
            if sheet_name and row and col:
                sheet = workbook[sheet_name]
                sheet.cell(row=row, column=col).value = field.get('edited_value', field.get('value', ''))
        
        # Salvar o arquivo atualizado
        workbook.save(output_path)
    
    def _save_docx_content(self, edited_data, output_path):
        """
        Salva conteúdo editado em formato Word
        
        Args:
            edited_data (dict): Dados editados
            output_path (str): Caminho do arquivo de saída
        """
        # Abrir o documento Word original
        doc = docx.Document(self.full_path)
        
        # Atualizar parágrafos
        for field in edited_data.get('fields', []):
            if not field.get('is_field'):
                continue
                
            # Parágrafos
            if 'paragraph_index' in field:
                para_index = field.get('paragraph_index')
                if para_index < len(doc.paragraphs):
                    original_text = doc.paragraphs[para_index].text
                    edited_value = field.get('edited_value', field.get('value', ''))
                    
                    # Substituir o campo de preenchimento pelo valor editado
                    new_text = original_text.replace('_____', edited_value).replace('____', edited_value).replace('___', edited_value)
                    
                    # Limpar o parágrafo e adicionar o texto editado
                    p = doc.paragraphs[para_index]
                    p.clear()
                    p.add_run(new_text)
            
            # Células de tabela
            elif 'table_index' in field and 'row_index' in field and 'col_index' in field:
                t_index = field.get('table_index')
                r_index = field.get('row_index')
                c_index = field.get('col_index')
                
                if t_index < len(doc.tables):
                    table = doc.tables[t_index]
                    if r_index < len(table.rows):
                        row = table.rows[r_index]
                        if c_index < len(row.cells):
                            cell = row.cells[c_index]
                            
                            original_text = cell.text
                            edited_value = field.get('edited_value', field.get('value', ''))
                            
                            # Substituir o campo de preenchimento pelo valor editado
                            new_text = original_text.replace('_____', edited_value).replace('____', edited_value).replace('___', edited_value)
                            
                            # Atualizar o conteúdo da célula
                            cell.text = new_text
        
        # Salvar o documento editado
        doc.save(output_path)
    
    def _save_pdf_content(self, edited_data, output_path):
        """
        Salva conteúdo editado em formato PDF
        
        Args:
            edited_data (dict): Dados editados
            output_path (str): Caminho do arquivo de saída
        """
        # Processo difere conforme o tipo de campo no PDF
        reader = PyPDF2.PdfReader(self.full_path)
        writer = PyPDF2.PdfWriter()
        
        # Adicionar todas as páginas originais
        for i in range(len(reader.pages)):
            writer.add_page(reader.pages[i])
        
        # Verificar se há campos de formulário
        form_fields = reader.get_fields()
        if form_fields:
            # PDF tem campos de formulário, atualizar os valores
            update_fields = {}
            
            for field in edited_data.get('fields', []):
                if field.get('is_form_field'):
                    field_id = field.get('id')
                    edited_value = field.get('edited_value', field.get('value', ''))
                    
                    if field_id in form_fields:
                        update_fields[field_id] = edited_value
            
            # Aplicar atualizações nos campos de formulário
            if update_fields:
                for i in range(len(reader.pages)):
                    writer.update_page_form_field_values(i, update_fields)
        
        # Salvar o PDF atualizado
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)


def extract_fields_from_file(file_path, forms_dir=None):
    """
    Função auxiliar para extrair campos de um arquivo para preenchimento
    
    Args:
        file_path (str): Caminho relativo do arquivo
        forms_dir (str, optional): Diretório base dos formulários
        
    Returns:
        list: Lista de campos detectados no arquivo
    """
    converter = OnlineEditorConverter(file_path, forms_dir)
    content = converter.get_editable_content()
    
    if not content.get('success', False):
        return []
    
    return content.get('fields', [])